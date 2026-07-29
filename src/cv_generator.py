import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


PROFILE_FILE = "data/profile/profile.json"
OPTIMIZATION_FILE = "data/profile/cv_optimization.json"

OUTPUT_FILE = "output/Ahmed_Tailored_CV.docx"


def load_json(path):

    with open(path, "r", encoding="utf-8") as file:
        return json.load(file)


def add_title(document, text):

    p = document.add_paragraph()

    run = p.add_run(text)

    run.bold = True
    run.font.size = Pt(13)


def add_bullet(document, text):

    p = document.add_paragraph(
        text,
        style="List Bullet"
    )

    p.paragraph_format.space_after = Pt(2)


def generate_cv(profile, optimization):

    document = Document()

    section = document.sections[0]

    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


    # Font

    style = document.styles["Normal"]

    style.font.name = "Aptos"
    style.font.size = Pt(10.5)


    # Header

    name = document.add_heading(level=0)

    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name.add_run(
        profile.get(
            "name",
            "Ahmed Ashraf EL Sayed Abdelbary"
        )
    )


    contact = document.add_paragraph()

    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact.add_run(
        f"{profile.get('email','')} | "
        f"{profile.get('phone','')} | "
        f"{profile.get('linkedin','')} | "
        f"{profile.get('location','')}"
    )


    # Summary

    add_title(
        document,
        "Professional Summary"
    )


    summary = optimization.get(
        "summary_improvement"
    )


    if summary:
        document.add_paragraph(summary)

    else:
        document.add_paragraph(
            profile.get(
                "summary",
                ""
            )
        )


    # Skills

    add_title(
        document,
        "Core Skills"
    )


    for skill in profile.get(
        "skills",
        []
    ):

        add_bullet(
            document,
            skill
        )


    # Experience

    add_title(
        document,
        "Professional Experience"
    )


    for company in profile.get(
        "companies",
        []
    ):

        p = document.add_paragraph()

        r = p.add_run(
            company.get(
                "company_name",
                company.get("name","")
            )
        )

        r.bold = True


        p.add_run(
            "\n" +
            company.get(
                "role",
                ""
            )
        )


        # Add original responsibilities

        for responsibility in company.get(
            "responsibilities",
            []
        ):

            add_bullet(
                document,
                responsibility
            )


    # Education

    add_title(
        document,
        "Education"
    )


    for edu in profile.get(
        "education",
        []
    ):

        add_bullet(
            document,
            json.dumps(
                edu,
                ensure_ascii=False
            )
        )


    # Courses

    add_title(
        document,
        "Certifications & Courses"
    )


    for course in profile.get(
        "courses",
        []
    ):

        add_bullet(
            document,
            json.dumps(
                course,
                ensure_ascii=False
            )
        )


    # Languages

    add_title(
        document,
        "Languages"
    )


    for lang in profile.get(
        "languages",
        []
    ):

        add_bullet(
            document,
            json.dumps(
                lang,
                ensure_ascii=False
            )
        )


    # ATS Keywords

    keywords = optimization.get(
        "keywords_to_emphasize",
        []
    )


    if keywords:

        add_title(
            document,
            "Keywords"
        )


        for keyword in keywords:

            add_bullet(
                document,
                keyword
            )


    os.makedirs(
        "output",
        exist_ok=True
    )


    document.save(
        OUTPUT_FILE
    )


    print(
        f"CV saved successfully: {OUTPUT_FILE}"
    )



if __name__ == "__main__":

    profile = load_json(
        PROFILE_FILE
    )

    optimization = load_json(
        OPTIMIZATION_FILE
    )


    generate_cv(
        profile,
        optimization
    )