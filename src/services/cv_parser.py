import os

import PyPDF2
from docx import Document


# ============================================================
# PDF Extraction
# ============================================================

def extract_pdf_text(uploaded_file):

    text = ""

    reader = PyPDF2.PdfReader(
        uploaded_file
    )

    for page in reader.pages:

        content = page.extract_text()

        if content:

            text += content + "\n"

    return text.strip()


# ============================================================
# DOCX Extraction
# ============================================================

def extract_docx_text(uploaded_file):

    document = Document(
        uploaded_file
    )

    text_parts = []

    # --------------------------------------------------------
    # Paragraphs
    # --------------------------------------------------------

    for paragraph in document.paragraphs:

        content = paragraph.text.strip()

        if content:

            text_parts.append(
                content
            )

    # --------------------------------------------------------
    # Tables
    #
    # Some CVs use tables for formatting.
    # We extract their text even though we don't use tables
    # in the generated ATS CV.
    # --------------------------------------------------------

    for table in document.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                content = cell.text.strip()

                if content:

                    row_text.append(
                        content
                    )

            if row_text:

                text_parts.append(
                    " | ".join(row_text)
                )

    return "\n".join(
        text_parts
    ).strip()


# ============================================================
# Main Extraction Function
# ============================================================

def extract_text(uploaded_file):

    if uploaded_file is None:

        return ""

    filename = (
        uploaded_file.name or ""
    ).lower()

    # --------------------------------------------------------
    # PDF
    # --------------------------------------------------------

    if filename.endswith(".pdf"):

        return extract_pdf_text(
            uploaded_file
        )

    # --------------------------------------------------------
    # DOCX
    # --------------------------------------------------------

    if filename.endswith(".docx"):

        return extract_docx_text(
            uploaded_file
        )

    # --------------------------------------------------------
    # Unsupported
    # --------------------------------------------------------

    raise ValueError(
        "Unsupported CV format. Please upload a PDF or DOCX file."
    )