import hashlib
import os
import sys
import re

import streamlit as st

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# Project Root
# ============================================================

sys.path.append(
    os.path.abspath(
        os.path.join(
            os.path.dirname(__file__),
            "../../.."
        )
    )
)


# ============================================================
# Services
# ============================================================

from src.services.cv_parser import extract_text
from src.services.cv_optimizer import (
    optimize_cv,
    calculate_ats_score
)


# ============================================================
# Page Configuration
# ============================================================

st.set_page_config(

    page_title="CV Builder",

    page_icon="📄",

    layout="wide"

)


# ============================================================
# Header
# ============================================================

st.title(
    "📄 CV Builder"
)

st.write(
    "Create an ATS-friendly CV tailored to your target role."
)


# ============================================================
# Session State
# ============================================================

defaults = {

    "cv_builder_text": "",

    "cv_builder_original_text": "",

    "cv_builder_result": None,

    "cv_builder_messages": [],

    "cv_builder_file_id": None,

    "current_ats_score": None

}


for key, value in defaults.items():

    if key not in st.session_state:

        st.session_state[key] = value


# ============================================================
# Upload CV
# ============================================================

uploaded_file = st.file_uploader(

    "Upload Your CV",

    type=[
        "pdf",
        "docx"
    ],

    key="cv_builder_upload"

)


if uploaded_file:

    file_bytes = uploaded_file.getvalue()

    file_id = hashlib.md5(
        file_bytes
    ).hexdigest()

    # --------------------------------------------------------
    # Detect New CV
    # --------------------------------------------------------

    if (
        st.session_state.cv_builder_file_id
        != file_id
    ):

        with st.spinner(
            "Analyzing Your CV..."
        ):

            extracted_text = extract_text(
                uploaded_file
            )

        st.session_state.cv_builder_original_text = (
            extracted_text
        )

        st.session_state.cv_builder_text = (
            extracted_text
        )

        st.session_state.cv_builder_result = None

        st.session_state.cv_builder_messages = []

        st.session_state.cv_builder_file_id = file_id

        st.session_state.current_ats_score = None

    st.success(
        "CV Successfully Loaded."
    )


# ============================================================
# Stop Until CV Is Uploaded
# ============================================================

if not st.session_state.cv_builder_text:

    st.info(
        "Upload Your CV To Start Building Your ATS-Friendly Version."
    )

    st.stop()


# ============================================================
# Target Role
# ============================================================

st.subheader(
    "🎯 Target Role"
)

target_job = st.text_input(

    "Target Job Title",

    placeholder=(
        "Example: Senior Demand Planning Manager "
        "(Optional)"
    )

)


job_description = st.text_area(

    "Job Description",

    height=220,

    placeholder=(
        "Paste The Job Description Here For Better ATS Optimization. "
        "You Can Leave This Empty If You Do Not Have One."
    )

)


# ============================================================
# User Instructions
# ============================================================

st.subheader(
    "✏️ What Would You Like To Change?"
)

st.info(
    """
Tell CareerAgent What You Want To Improve In Natural Language.

You Can Ask For Things Such As:

• Rewrite My Samsung Experience To Sound More Strategic.

• Rewrite My Experience At Company X And Highlight The Work I Did.

• Add A Professional Skills Section Using The Skills Already Shown In My CV.

• Make My Demand Planning Experience Stronger.

• Make My CV More Suitable For Supply Chain Manager Roles.

• Rewrite My Achievements Using Stronger Action Verbs.

• Improve My Professional Summary.

• Make My Experience More Concise.

You Can Write Your Request Naturally. CareerAgent Will Analyze
Your CV And Your Instructions Before Making The Changes.

Important: CareerAgent Will Not Invent Experience Or Skills.
Anything That Requires Your Confirmation Will Be Flagged.
"""
)

user_instructions = st.text_area(

    "Your Instructions",

    height=180,

    placeholder=(
        "Example: In My Samsung Experience, I Managed Demand And "
        "Supply Planning Across MENA. Rewrite This More Professionally "
        "And Highlight My Forecasting And Stakeholder Management."
    )

)


# ============================================================
# ATS Score
# ============================================================

st.divider()

st.subheader(
    "📊 Current ATS Score"
)

if st.button(
    "🔍 Check Current ATS Score",
    use_container_width=True
):

    with st.spinner(
        "Analyzing ATS Compatibility..."
    ):

        score = calculate_ats_score(

            st.session_state.cv_builder_text,

            target_job,

            job_description

        )

    st.session_state.current_ats_score = score


if (
    st.session_state.current_ats_score
    is not None
):

    st.metric(

        "Current ATS Score",

        f"{st.session_state.current_ats_score}%"

    )


# ============================================================
# Generate
# ============================================================

st.divider()

generate_button = st.button(

    "✨ Generate ATS-Friendly CV",

    use_container_width=True,

    type="primary"

)


if generate_button:

    with st.spinner(
        "AI Is Optimizing Your CV..."
    ):

        result = optimize_cv(

            st.session_state.cv_builder_text,

            target_job,

            job_description,

            user_instructions,

            str(
                st.session_state.cv_builder_messages
            )

        )

    st.session_state.cv_builder_result = result

    st.session_state.cv_builder_text = (
        result["cv_text"]
    )

    st.session_state.cv_builder_messages.append({

        "role": "user",

        "content":
            user_instructions or
            "Generate An ATS-Friendly CV."

    })

    st.session_state.cv_builder_messages.append({

        "role": "assistant",

        "content":
            result["summary"]

    })

    st.rerun()


# ============================================================
# Results
# ============================================================

result = st.session_state.cv_builder_result


if result:

    st.divider()

    st.subheader(
        "📈 ATS Score Improvement"
    )

    col1, col2, col3 = st.columns(3)

    with col1:

        st.metric(

            "Before",

            f"{result['before_score']}%"

        )

    with col2:

        st.metric(

            "After",

            f"{result['after_score']}%"

        )

    with col3:

        change = result["score_change"]

        st.metric(

            "Improvement",

            f"{change:+d}%"

        )


    # ========================================================
    # AI Summary
    # ========================================================

    if result.get(
        "summary"
    ):

        st.info(
            result["summary"]
        )


    # ========================================================
    # Changes
    # ========================================================

    changes = result.get(
        "changes",
        []
    )

    if changes:

        with st.expander(
            "🔍 Changes Made",
            expanded=True
        ):

            for change in changes:

                st.write(
                    f"• {change}"
                )


    # ========================================================
    # Warnings
    # ========================================================

    warnings = result.get(
        "warnings",
        []
    )

    if warnings:

        with st.expander(
            "⚠️ Items Requiring Confirmation"
        ):

            for warning in warnings:

                st.warning(
                    warning
                )


    # ========================================================
    # Updated CV
    # ========================================================

    st.subheader(
        "📄 Updated CV"
    )

    st.text_area(

        "Generated CV",

        value=st.session_state.cv_builder_text,

        height=700,

        key="generated_cv_display"

    )


    # ========================================================
    # Professional DOCX Generator
    # ========================================================

    def create_professional_docx(
        cv_text
    ):

        document = Document()

        # ----------------------------------------------------
        # Page Setup
        # ----------------------------------------------------

        section = document.sections[0]

        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(45)
        section.right_margin = Pt(45)

        # ----------------------------------------------------
        # Default Font
        # ----------------------------------------------------

        styles = document.styles

        normal = styles["Normal"]

        normal.font.name = "Arial"
        normal.font.size = Pt(10)

        # ----------------------------------------------------
        # Recognized Section Headers
        # ----------------------------------------------------

        section_headers = {

            "professional summary",
            "summary",
            "profile",
            "professional experience",
            "work experience",
            "experience",
            "education",
            "skills",
            "core skills",
            "technical skills",
            "certifications",
            "projects",
            "additional information",
            "languages",
            "awards",
            "achievements"

        }

        lines = cv_text.splitlines()

        first_content = True

        for raw_line in lines:

            line = raw_line.strip()

            if not line:

                continue

            clean_line = re.sub(
                r"^[•●▪◦\-–—]\s*",
                "",
                line
            ).strip()

            normalized = clean_line.lower()

            # ------------------------------------------------
            # Name / First Content
            # ------------------------------------------------

            if first_content:

                paragraph = document.add_paragraph()

                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                )

                run = paragraph.add_run(
                    clean_line
                )

                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(18)

                first_content = False

                continue

            # ------------------------------------------------
            # Section Heading
            # ------------------------------------------------

            if normalized in section_headers:

                paragraph = document.add_paragraph()

                paragraph.paragraph_format.space_before = Pt(10)
                paragraph.paragraph_format.space_after = Pt(4)

                run = paragraph.add_run(
                    clean_line.upper()
                )

                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(11)

                continue

            # ------------------------------------------------
            # Bullet
            # ------------------------------------------------

            if (
                raw_line.startswith("•")
                or raw_line.startswith("-")
                or raw_line.startswith("●")
                or raw_line.startswith("▪")
                or raw_line.startswith("◦")
            ):

                paragraph = document.add_paragraph(
                    style="List Bullet"
                )

                paragraph.paragraph_format.space_after = Pt(2)

                run = paragraph.add_run(
                    clean_line
                )

                run.font.name = "Arial"
                run.font.size = Pt(10)

                continue

            # ------------------------------------------------
            # Normal Text
            # ------------------------------------------------

            paragraph = document.add_paragraph()

            paragraph.paragraph_format.space_after = Pt(3)

            run = paragraph.add_run(
                clean_line
            )

            run.font.name = "Arial"
            run.font.size = Pt(10)

        return document


    # ========================================================
    # Download DOCX
    # ========================================================

    document = create_professional_docx(

        st.session_state.cv_builder_text

    )

    output_path = os.path.join(

        os.getcwd(),

        "CareerAgent_ATS_CV.docx"

    )

    document.save(
        output_path
    )

    with open(
        output_path,
        "rb"
    ) as file:

        st.download_button(

            "⬇️ Download ATS-Friendly CV",

            data=file,

            file_name="CareerAgent_ATS_CV.docx",

            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),

            use_container_width=True

        )


    # ========================================================
    # Conversation
    # ========================================================

    st.divider()

    st.subheader(
        "💬 Continue Editing Your CV"
    )

    st.write(
        "Review The CV Above And Tell CareerAgent What You Want "
        "To Change Next."
    )


    # --------------------------------------------------------
    # Previous Conversation
    # --------------------------------------------------------

    if st.session_state.cv_builder_messages:

        with st.expander(
            "💬 Edit History",
            expanded=True
        ):

            for message in (
                st.session_state.cv_builder_messages
            ):

                if message["role"] == "user":

                    st.markdown(
                        f"**You:** {message['content']}"
                    )

                else:

                    st.markdown(
                        f"**CareerAgent:** {message['content']}"
                    )


    # --------------------------------------------------------
    # Follow-up
    # --------------------------------------------------------

    follow_up = st.text_area(

        "Your Next Edit",

        height=120,

        placeholder=(
            "Example: Make My Samsung Experience More Senior, "
            "but keep the same facts. Also improve the summary."
        ),

        key="cv_follow_up"

    )


    if st.button(
        "✏️ Apply This Edit",
        use_container_width=True
    ):

        if not follow_up.strip():

            st.warning(
                "Please Enter An Edit Before Continuing."
            )

        else:

            with st.spinner(
                "Applying Your Edit..."
            ):

                updated_result = optimize_cv(

                    st.session_state.cv_builder_text,

                    target_job,

                    job_description,

                    follow_up,

                    str(
                        st.session_state.cv_builder_messages
                    )

                )

            st.session_state.cv_builder_result = (
                updated_result
            )

            st.session_state.cv_builder_text = (
                updated_result["cv_text"]
            )

            st.session_state.cv_builder_messages.append({

                "role": "user",

                "content":
                    follow_up

            })

            st.session_state.cv_builder_messages.append({

                "role": "assistant",

                "content":
                    updated_result["summary"]

            })

            st.rerun()