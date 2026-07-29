import json
import requests
import os
from docx import Document
from docx.shared import Pt


PROFILE_FILE = "data/profile/profile.json"
JOB_FILE = "data/profile/job_profile.json"
MATCH_FILE = "data/profile/match_report.json"

OUTPUT_FILE = "output/CV_Optimization_Report.docx"


def load_json(path):

    with open(path, "r", encoding="utf-8") as file:
        return json.load(file)


def ask_ai(prompt):

    print("Sending analysis request to AI...")

    response = requests.post(
        "http://localhost:11434/api/generate",
        json={
            "model": "qwen2.5:1.5b",
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": 0,
                "num_ctx": 2048
            }
        },
        timeout=300
    )

    response.raise_for_status()

    result = response.json().get(
        "response",
        ""
    )

    if not result:
        raise Exception(
            "AI returned empty response"
        )

    return result


def create_advice(profile, job, match):

    prompt = f"""
You are a senior ATS CV consultant.

Analyze the candidate CV against the job description.

Your task is to provide editing advice ONLY.

STRICT RULES:

- Do NOT rewrite the CV.
- Do NOT remove existing relevant skills.
- Do NOT remove experience.
- Do NOT remove achievements.
- Do NOT invent experience.
- Do NOT add company-specific experience.
- Do NOT add Amazon, Google, Microsoft, or any company-specific keywords unless the candidate worked there.
- Do NOT suggest fake tools or systems.
- Keep the candidate's career identity unchanged.

Only recommend:
- Better wording of existing bullets.
- Keywords that are realistic for the candidate.
- Where to place keywords.
- Skills ordering improvements.
- Summary improvements.

Your recommendations must be realistic and transferable.

Create the report using this structure:

1. Overall Assessment

2. Professional Summary Improvements

Include:
- Current issue
- Suggested direction

3. Skills Section Changes

Include:
- Add
- Reorder
- Avoid removing relevant existing skills

4. Experience Section Improvements

For each recommendation include:

Company:
Current Bullet:
Suggested Improvement:
Reason:

5. Missing ATS Keywords

Only include keywords that match the candidate's actual experience.

6. Final ATS Checklist


Candidate CV:

{json.dumps(profile, indent=2, ensure_ascii=False)}


Job Description:

{json.dumps(job, indent=2, ensure_ascii=False)}


Current Match Report:

{json.dumps(match, indent=2, ensure_ascii=False)}

"""

    return ask_ai(prompt)


def create_word_report(text):

    os.makedirs(
        "output",
        exist_ok=True
    )

    document = Document()


    style = document.styles["Normal"]

    style.font.name = "Aptos"
    style.font.size = Pt(11)


    document.add_heading(
        "CV Optimization Report",
        level=1
    )


    document.add_paragraph(
        text
    )


    document.save(
        OUTPUT_FILE
    )


    print(
        f"Report saved: {OUTPUT_FILE}"
    )


if __name__ == "__main__":

    print(
        "Creating CV advice report..."
    )


    profile = load_json(
        PROFILE_FILE
    )

    job = load_json(
        JOB_FILE
    )

    match = load_json(
        MATCH_FILE
    )


    report = create_advice(
        profile,
        job,
        match
    )


    print("\nREPORT:\n")

    print(report)


    create_word_report(
        report
    )